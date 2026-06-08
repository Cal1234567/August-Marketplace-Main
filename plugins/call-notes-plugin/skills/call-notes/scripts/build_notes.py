#!/usr/bin/env python3
"""
build_notes.py — Render a call-notes briefing into the August Group house Word template.

It clones the bundled "Notes Template.docx" (preserving the logo header, footer,
margins, theme and named paragraph styles), removes the placeholder paragraphs, and
lays in the supplied content as styled paragraphs *before* the document's section
properties (so headers/footers stay intact).

Content is supplied as a JSON spec describing an ordered list of blocks. Each block is
one paragraph. python-docx writes proper UTF-8 and XML-escapes text automatically, so
smart quotes / em dashes / arrows come out clean (no "â€"" mojibake).

Note: the builder replaces the entire document *body*. The template's chrome (logo
header, footer, margins, styles) lives outside the body and is preserved, but this
assumes the template body holds only placeholder paragraphs — any real body content in
the template would be removed.

Usage
-----
    python build_notes.py SPEC.json [--out OUTPUT.docx] [--template TEMPLATE.docx]
    python build_notes.py --dump-styles            # list style ids/names in the template

Paths: prefer forward slashes everywhere (they work on Windows and avoid JSON backslash
escaping). To override the output without touching the spec, pass --out — a normal
Windows path is fine there since command-line args aren't JSON.

Spec format
-----------
{
  "output_path": "C:/Users/.../<Call name> - Notes.docx",
  "blocks": [
    {"style": "ATitle",     "text": "City Golf — Investor Intro Call"},
    {"style": "BSubheader", "text": "Snapshot"},
    {"style": "FBullet", "label": "Round: ", "text": "$15M at $30M pre / $45M post."},
    {"style": "FBullet", "level": 1, "text": "$12M already committed; ~$2–3M left for strategics."},
    {"style": "FBullet", "key": true, "text": "Investment is into the whole US holdco, not a single site."},
    {"style": "FBullet", "runs": [
        {"text": "Payback "}, {"text": "~2 years at 65% utilization", "bold": true, "highlight": true},
        {"text": "; breakeven at 25%."}]}
  ]
}

Block fields (only these keys are allowed)
-------------------------------------------
  style  (required)  one of: ATitle, BSubheader, CSubheader, DSubheader,
                     EParagraph, FBullet, Risks, RisksHeader
  label  (optional)  bold lead-in run, written verbatim incl. trailing punctuation/space
  text   (optional)  the normal-weight run that follows the label
  level  (optional)  FBullet sub-bullet depth: 0 (default) → •, 1 → o, 2 → ▪, 3 → •
  key    (optional)  true = critical takeaway: bold + yellow-highlight the whole bullet
  runs   (optional)  [{text, bold?, italic?, highlight?}] for fine control (e.g. highlight
                     only a phrase). When present, it replaces label/text. highlight may be
                     true (= yellow) or a color name (yellow/green/turquoise/pink/gray).
  (a block needs at least one of label / text / runs-with-text)
"""

import argparse
import json
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Twips
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    sys.exit("python-docx is required. Install with:  python -m pip install python-docx")

ALLOWED_STYLES = {
    "ATitle", "BSubheader", "CSubheader", "DSubheader",
    "EParagraph", "FBullet", "Risks", "RisksHeader",
}
ALLOWED_BLOCK_KEYS = {"style", "label", "text", "level", "key", "runs"}
ALLOWED_RUN_KEYS = {"text", "bold", "italic", "highlight"}

FBULLET_NUMID = 1   # numbering id used by the "F) Bullet" style (a 4-level bullet list)
MAX_LEVEL = 3
# Left indent (twips) per sub-bullet level, matching the template's numbering levels.
LEVEL_INDENT = {1: 1080, 2: 1800, 3: 2520}

_HIGHLIGHTS = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "bright_green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "turquoise": WD_COLOR_INDEX.TURQUOISE,
    "cyan": WD_COLOR_INDEX.TURQUOISE,
    "pink": WD_COLOR_INDEX.PINK,
    "magenta": WD_COLOR_INDEX.PINK,
    "gray": WD_COLOR_INDEX.GRAY_25,
    "grey": WD_COLOR_INDEX.GRAY_25,
}


def _highlight(value):
    if isinstance(value, str):
        return _HIGHLIGHTS.get(value.strip().lower(), WD_COLOR_INDEX.YELLOW)
    return WD_COLOR_INDEX.YELLOW  # True or anything truthy → yellow


def default_template() -> Path:
    return Path(__file__).resolve().parent.parent / "assets" / "Notes Template.docx"


def styles_by_id(document):
    return {s.style_id: s for s in document.styles}


def dump_styles(template: Path) -> None:
    document = Document(str(template))
    print(f"Styles in {template.name}:")
    for s in document.styles:
        print(f"  {(s.style_id or ''):<20} {(s.name or '')!r}")


def set_bullet_level(para, level):
    """Override the paragraph's list level so it renders as a sub-bullet (o / ▪)."""
    numPr = para._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = level
    numPr.get_or_add_numId().val = FBULLET_NUMID
    if level in LEVEL_INDENT:
        para.paragraph_format.left_indent = Twips(LEVEL_INDENT[level])


def validate(blocks, sid_map):
    errors, warnings = [], []
    if not isinstance(blocks, list) or not blocks:
        return (["Spec 'blocks' must be a non-empty list."], warnings)

    for i, b in enumerate(blocks):
        if not isinstance(b, dict):
            errors.append(f"block {i}: must be an object, got {type(b).__name__}")
            continue
        unknown = set(b) - ALLOWED_BLOCK_KEYS
        if unknown:
            errors.append(f"block {i}: unknown key(s) {sorted(unknown)} — allowed: {sorted(ALLOWED_BLOCK_KEYS)}")
        style = b.get("style")
        if not style:
            errors.append(f"block {i}: missing 'style'")
        elif style not in ALLOWED_STYLES:
            errors.append(f"block {i}: unknown style {style!r} — allowed: {', '.join(sorted(ALLOWED_STYLES))}")
        elif style not in sid_map:
            errors.append(f"block {i}: style {style!r} not found in this template")

        runs = b.get("runs")
        has_text = bool(b.get("label") or b.get("text"))
        if runs is not None:
            if not isinstance(runs, list) or not runs:
                errors.append(f"block {i}: 'runs' must be a non-empty list")
            else:
                for j, r in enumerate(runs):
                    if not isinstance(r, dict):
                        errors.append(f"block {i} run {j}: must be an object")
                        continue
                    bad = set(r) - ALLOWED_RUN_KEYS
                    if bad:
                        errors.append(f"block {i} run {j}: unknown key(s) {sorted(bad)} — allowed: {sorted(ALLOWED_RUN_KEYS)}")
                has_text = has_text or any(isinstance(r, dict) and r.get("text") for r in runs)
        if not has_text:
            warnings.append(f"block {i} ({style or '?'}): empty — will render a blank paragraph")

        level = b.get("level")
        if level is not None:
            if not isinstance(level, int) or isinstance(level, bool) or level < 0:
                errors.append(f"block {i}: 'level' must be an integer ≥ 0")
            elif level > 0 and style != "FBullet":
                warnings.append(f"block {i}: 'level' only applies to FBullet — ignored for {style}")
        if b.get("key") is not None and not isinstance(b.get("key"), bool):
            errors.append(f"block {i}: 'key' must be true/false")

    titles = sum(1 for b in blocks if isinstance(b, dict) and b.get("style") == "ATitle")
    if titles == 0:
        warnings.append("no ATitle block — the document will have no title")
    elif titles > 1:
        warnings.append(f"{titles} ATitle blocks — expected exactly one")
    return (errors, warnings)


def render_block(para, b):
    """Add runs to an already-styled paragraph. Style + bullet level are set by the caller."""
    runs = b.get("runs")
    if runs:
        for r in runs:
            run = para.add_run(r.get("text", ""))
            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            if r.get("highlight"):
                run.font.highlight_color = _highlight(r["highlight"])
        return

    key = bool(b.get("key"))
    label, text = b.get("label"), b.get("text")
    if label:
        run = para.add_run(label)
        run.bold = True
        if key:
            run.font.highlight_color = _highlight(True)
    if text:
        run = para.add_run(text)
        if key:
            run.bold = True
            run.font.highlight_color = _highlight(True)


def _unlink_quiet(p: Path) -> None:
    try:
        p.unlink()
    except OSError:
        pass


def save_atomic(document, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = out_path.parent / (out_path.name + ".tmp")
    try:
        document.save(str(tmp))
    except PermissionError:
        _unlink_quiet(tmp)
        sys.exit(f"Permission denied writing to:\n  {out_path.parent}\nIs the folder writable?")
    try:
        os.replace(str(tmp), str(out_path))
    except PermissionError:
        _unlink_quiet(tmp)
        sys.exit(
            f"Could not write the output file:\n  {out_path}\n"
            "It is almost certainly open in Word — close it and re-run."
        )


def build(spec: dict, out_path: Path, template: Path) -> Path:
    if not isinstance(spec, dict):
        sys.exit("Spec must be a JSON object with a 'blocks' list.")

    document = Document(str(template))
    sid_map = styles_by_id(document)
    blocks = spec.get("blocks", [])

    errors, warnings = validate(blocks, sid_map)
    for w in warnings:
        print(f"warning: {w}", file=sys.stderr)
    if errors:
        sys.exit("Spec validation failed:\n  - " + "\n  - ".join(errors))

    placeholders = list(document.paragraphs)
    if not placeholders:
        sys.exit("Template has no body paragraphs to anchor against.")
    anchor = placeholders[0]

    for b in blocks:
        para = anchor.insert_paragraph_before()
        para.style = sid_map[b["style"]]
        level = int(b.get("level") or 0)
        if level > 0 and b["style"] == "FBullet":
            set_bullet_level(para, min(level, MAX_LEVEL))
        render_block(para, b)

    for p in placeholders:
        p._element.getparent().remove(p._element)

    save_atomic(document, out_path)
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser(description="Render call notes into the August Group Word template.")
    ap.add_argument("spec", nargs="?", help="Path to the JSON spec.")
    ap.add_argument("--out", help="Output .docx path (overrides spec.output_path). A normal Windows path is fine here.")
    ap.add_argument("--template", help="Override the bundled template path.")
    ap.add_argument("--dump-styles", action="store_true", help="List template styles and exit.")
    args = ap.parse_args()

    template = Path(args.template) if args.template else default_template()
    if not template.exists():
        sys.exit(f"Template not found: {template}")

    if args.dump_styles:
        dump_styles(template)
        return

    if not args.spec:
        ap.error("a spec.json path is required (or use --dump-styles)")

    spec_path = Path(args.spec)
    if not spec_path.exists():
        sys.exit(f"Spec not found: {spec_path}")
    try:
        spec = json.loads(spec_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        sys.exit(
            f"Spec is not valid JSON ({e}).\n"
            "Tip: use forward slashes in paths (e.g. C:/Users/...) — backslashes are invalid JSON escapes."
        )

    out = args.out or spec.get("output_path")
    if not out:
        sys.exit("No output path: pass --out or set 'output_path' in the spec.")
    out_path = Path(out)
    if out_path.suffix.lower() != ".docx":
        print(f"warning: output path does not end in .docx: {out_path}", file=sys.stderr)

    saved = build(spec, out_path, template)
    print(f"Saved: {saved}")


if __name__ == "__main__":
    main()
