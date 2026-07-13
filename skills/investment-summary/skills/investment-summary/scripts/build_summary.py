#!/usr/bin/env python3
"""
build_summary.py — Render an August Group Investment Summary (.docx) in the house style:
a clean titled document (logo header -> title -> sections), teal-headed data tables, and the
standard section structure ending Considerations -> Appendix -> Disclaimer. This matches the
golden HARVEST Clean Eats summary.

The Addepar/Form/Radiant deck cover (hero banner -> gold "INVESTMENT SUMMARY" label ->
4-column facts table) is OPT-IN: set "cover": true in the spec to enable it.

Built on the Sisyphus golden template (five named styles + logo header + page-number footer).

USAGE
    python build_summary.py spec.json
    python build_summary.py spec.json --out "C:/path/Acme - Investment Summary.docx"

SPEC (JSON)
{
  "title": "ACME ROBOTICS",
  "filename": "Acme Robotics - Investment Summary.docx",   # optional
  "output_dir": "C:/Users/.../Desktop",                    # optional (default = cwd)

  "cover": false,                                          # opt-in deck cover (banner+label+facts)
  "banner_image": "C:/path/acme_logo_white.png",           # cover only; else a teal band w/ the title
  "facts": {                                               # cover only: 4-column facts table
      "asset_class": "Direct Investment – Growth Equity",
      "risk_level": "High",
      "minimum_investment": "$100k USD",
      "liquidity": "Illiquid"
  },
  "canadian": false,                                        # true => "For Canadian residents only" note

  "blocks": [ ...ordered body content... ]
}

BLOCK TYPES
  {"type":"h1","text":"Executive Summary"}      section heading -> . Subheader (Summary)
  {"type":"h2","text":"Core Product"}           sub-heading     -> . Subheading2 (gold)
  {"type":"p","text":"Acme (\"Acme\"...)..."}   body paragraph  -> . Paragraph (Summary)
  {"type":"bullets","items":[{"text":"**Lead-in**: ...","level":0},{"text":"detail","level":1}]}
  {"type":"table","headers":["","2024A","2025E"],          native Word table (teal header, white text)
      "rows":[["ARR","$247M","$310M"],["Total","$247M","$310M"]],
      "emphasis_rows":[1],                                 # 0-based rows to shade/bold (totals)
      "source":"Source: Acme Investor Model"}
  {"type":"source","text":"Source: ..."}        gold italic source line (for pasted images/charts)
  {"type":"risk_factors"}                        verbatim Risk Factors section (opt-in; golden omits it)
  {"type":"disclaimer","variant":"short"|"long"} verbatim Disclaimer (default = AGC standard; golden's)

INLINE MARKUP in any "text": **bold**, *italic* (use for "(see Exhibit 1)" refs), and
  [[gap-flag]] -> renders bold dark-red on a yellow highlight. Use the gap-flag for anything
  you could NOT verify (a tenure, date, school, or unconfirmed identity), e.g.
  "...at Loblaw [[NEED: confirm tenure]]" — a visible to-do for the human finishing the doc.

Complex charts / positioning maps / multi-panel financial exhibits should be built in
Excel/PowerPoint and pasted into the .docx as pictures, each with a {"type":"source"} line.
"""
import sys, os, re, json, argparse

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx is required.  Install with:  pip install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "..", "assets", "Golden_Template.docx")

# --- Named styles (exact UI names in the template) -------------------------------
S_TITLE = ". Title (Summary)"
S_H1    = ". Subheader (Summary)"
S_H2    = ". Subheading2"
S_BODY  = ". Paragraph (Summary)"

# --- Brand palette ---------------------------------------------------------------
TEAL     = "1C4849"   # hero banner + data-table header fill
GOLD_HEX = "B29659"   # accent3 (label, source lines, facts-table rule)
NAVY_HEX = "000825"   # dk2
GRAYFILL = "D9D9D9"   # facts-table fill
EMPH     = "DCE4E4"   # pale-teal emphasis-row fill
GOLD     = RGBColor(0xB2, 0x96, 0x59)
NAVY     = RGBColor(0x00, 0x08, 0x25)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BLACK    = RGBColor(0x00, 0x00, 0x00)
FLAG_RED = RGBColor(0xC0, 0x00, 0x00)   # gap-flag text ([[NEED: ...]] / [[VERIFY: ...]])

CONTENT_W = 7.5       # inches, with 0.5" margins on US Letter
BULLET_NUMID = 32     # template's clean  ● / o  bullet hierarchy

# --- Verbatim boilerplate (operating-company variants) ---------------------------
RISK_FACTORS = [
    "Growth equity capital investing involves a high degree of business and financial risk that can result in substantial losses.",
    "An investment in this SPV/fund is highly speculative, involves a high degree of risk and could result in the loss of part or "
    "all of a Limited Partner's capital contribution. Therefore, Limited Partners should not subscribe for LP Interests unless they "
    "can bear such a loss. Moreover, there can be no assurance that the fund's investment objectives will be achieved, and investment "
    "results may vary materially from one reporting period to the next. Consequently, an investment in the fund is suitable only for "
    "sophisticated investors with substantial other assets who are capable of making an informed independent decision as to the risks "
    "involved in an investment in the fund.",
]
DISCLAIMER_SHORT = [
    "The information contained in this material is subject to change without notice and August Group Capital will not be held liable "
    "for any inaccuracies or misprints. The information within this Investment Summary is meant for informational purposes and is not "
    "meant to elicit an individual investment. Information is taken from sources believed to be accurate.",
    "August Group Capital is a registered Portfolio Manager with the Autorité des Marchés Financiers (Quebec), Ontario Securities "
    "Commission, Alberta Securities Commission, British Columbia Securities Commission, the Manitoba Securities Commission, and a "
    "Registered Investment Advisor with the U.S. Securities and Exchange Commission.",
    "The SPV manager in question is not an affiliate of August Group Capital. More information may be obtained by contacting August "
    "Group Capital (the registrant).",
]
DISCLAIMER_LONG = [
    "This Investment Summary is intended to be a factual overview of an investment, provided to you for informational purposes. The "
    "information contained in this Investment Summary is derived from sources that we believe to be reliable and accurate, including "
    "information received from the SPV/Fund and/or companies related to the SPV/Fund. The August Group Capital Limited makes no "
    "representations or warranties as to accuracy, completeness or reliability of such information, and will not be held liable for any "
    "inaccuracies or misprints derived from such information.",
    "This Investment Summary is not intended to be, and should not be taken to be, an offer to sell or a solicitation of an offer to buy "
    "the investment that is the subject of the Investment Summary, or any other investment. Investing your money in investment products "
    "involves the risk of loss, and past performance of investment products is not a guarantee of future returns.",
    "This Investment Summary is confidential and intended solely for the use of the individual or entity to whom it is addressed. Any "
    "unauthorized review, use, disclosure, or distribution is prohibited. If you are not the intended recipient, please contact the "
    "sender and destroy any copies that may have been inadvertently sent to you.",
    "August Group Capital is a registered Portfolio Manager with the Autorité des Marchés Financiers (Quebec), Ontario Securities "
    "Commission, Alberta Securities Commission, British Columbia Securities Commission, the Manitoba Securities Commission, and a "
    "Registered Investment Advisor with the U.S. Securities and Exchange Commission.",
]
# The standard, current house disclaimer (AGC-attributed) — used by the golden HARVEST
# summary. This is the DEFAULT; "short"/"long" remain available via the block's "variant".
# Bullets intentionally carry NO trailing period, matching the golden document exactly.
DISCLAIMER_STANDARD = [
    "The information herein was prepared by The August Group Capital Limited (\"AGC\") and is derived from sources AGC believes to be "
    "reliable, including information received from the fund manager or company management. AGC makes no representations or warranties "
    "as to accuracy, completeness or reliability of such information, and will not be held liable for any inaccuracies or misprints "
    "derived from such information",
    "This report is not intended to be, and should not be taken to be, an offer to sell or a solicitation of an offer to buy the "
    "investment or investments that are the subject of the report, or any other investment. Investing your money in investment products "
    "involves the risk of loss, and past performance of investment products is not a guarantee of future returns. Investments in "
    "private funds are subject to a number of risks, including liquidity risk and the risk of loss of capital, and are designed for "
    "investors who understand and are able to withstand the risks",
    "This report is confidential and intended solely for the use of the individual or entity to whom it is addressed. Any unauthorized "
    "review, use, disclosure, or distribution is prohibited. If you are not the intended recipient, please contact the sender and "
    "destroy any copies that may have been inadvertently sent to you",
    "AGC is registered as a portfolio manager in Alberta, British Columbia, Manitoba, Ontario, and Quebec, and is registered as an "
    "investment adviser in the United States. The investment or investments that are the subject of this report are not affiliated with AGC",
]

# --- low-level helpers -----------------------------------------------------------
_INLINE = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|(\[\[.+?\]\])')

def _emit(paragraph, text, bold=False, italic=False, flag=False):
    """Emit text into runs, turning '\\n' into real line breaks.
    flag=True renders a gap marker ([[...]]) as bold dark-red on a yellow highlight,
    so a human collaborator can see at a glance exactly what still needs filling in."""
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i:
            paragraph.add_run().add_break()
        r = paragraph.add_run(part)
        if bold: r.bold = True
        if italic: r.italic = True
        if flag:
            r.bold = True
            r.font.color.rgb = FLAG_RED
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW

def add_runs(paragraph, text):
    """Add runs honoring **bold**, *italic*, and [[gap-flag]] inline markup and '\\n' breaks."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _emit(paragraph, text[pos:m.start()])
        if m.group(1) is not None:
            _emit(paragraph, m.group(1), bold=True)
        elif m.group(2) is not None:
            _emit(paragraph, m.group(2), italic=True)
        else:
            _emit(paragraph, m.group(3), flag=True)   # render [[...]] literally, highlighted
        pos = m.end()
    if pos < len(text):
        _emit(paragraph, text[pos:])
    return paragraph

def set_bullet(paragraph, level=0, num_id=BULLET_NUMID):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl');  ilvl.set(qn('w:val'), str(int(level)))
    nid  = OxmlElement('w:numId'); nid.set(qn('w:val'), str(int(num_id)))
    numPr.append(ilvl); numPr.append(nid)
    pPr.append(numPr)

def shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def cell_border(cell, edges, hex_color="auto", sz="4", val="single"):
    """edges: subset of ('top','bottom','left','right'); sz in eighths of a point."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = tcPr.find(qn('w:tcBorders'))
    if tcB is None:
        tcB = OxmlElement('w:tcBorders'); tcPr.append(tcB)
    for edge in edges:
        el = tcB.find(qn('w:' + edge))
        if el is None:
            el = OxmlElement('w:' + edge); tcB.append(el)
        el.set(qn('w:val'), val); el.set(qn('w:sz'), sz); el.set(qn('w:space'), '0'); el.set(qn('w:color'), hex_color)

def set_table_grid(table, width_in=CONTENT_W):
    """Fixed layout at a known total width."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    for tag in ('w:tblW', 'w:tblLayout'):
        ex = tblPr.find(qn(tag))
        if ex is not None:
            tblPr.remove(ex)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:type'), 'dxa'); tblW.set(qn('w:w'), str(int(width_in * 1440)))
    tblPr.append(tblW)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)

def fmt_cell(cell, text, *, color=BLACK, bold=False, size=None, align="left", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    add_runs(p, text)
    for r in p.runs:
        r.font.name = "Arial"
        if size: r.font.size = Pt(size)
        if bold: r.bold = True
        if italic: r.italic = True
        r.font.color.rgb = color

# --- cover -----------------------------------------------------------------------
def add_banner(doc, title, banner_image=None):
    if banner_image and os.path.exists(banner_image):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(6)
        p.add_run().add_picture(banner_image, width=Inches(CONTENT_W))
        return
    # teal band with the company name in white
    tbl = doc.add_table(rows=1, cols=1); set_table_grid(tbl)
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(CONTENT_W)
    shade(cell, TEAL)
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), '2100'); h.set(qn('w:hRule'), 'atLeast'); trPr.append(h)
    cell.vertical_alignment = 1  # center
    fmt_cell(cell, title, color=WHITE, bold=True, size=30, align="center")

def add_label(doc, text="INVESTMENT SUMMARY"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = GOLD
    rPr = r._r.get_or_add_rPr()
    sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), '40'); rPr.append(sp)  # letter-spacing

def add_facts_table(doc, facts):
    labels = [("Asset Class", facts.get("asset_class", "Direct Investment – Growth Equity")),
              ("Risk Level", facts.get("risk_level", "High")),
              ("Minimum Investment", facts.get("minimum_investment", "$100k USD")),
              ("Liquidity", facts.get("liquidity", "Illiquid"))]
    tbl = doc.add_table(rows=2, cols=4); set_table_grid(tbl)
    for j, (head, val) in enumerate(labels):
        hc = tbl.rows[0].cells[j]; vc = tbl.rows[1].cells[j]
        hc.width = Inches(CONTENT_W / 4); vc.width = Inches(CONTENT_W / 4)
        shade(hc, GRAYFILL); shade(vc, GRAYFILL)
        fmt_cell(hc, head, color=NAVY, bold=True, size=10, align="center")
        fmt_cell(vc, val, color=RGBColor(0x40, 0x40, 0x40), size=10, align="center")
        cell_border(hc, ("bottom",), hex_color=GOLD_HEX, sz="12")  # gold rule under header
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# --- data table block ------------------------------------------------------------
def add_data_table(doc, blk):
    headers = blk.get("headers", [])
    rows = blk.get("rows", [])
    emph = set(blk.get("emphasis_rows", []))
    ncol = max([len(headers)] + [len(r) for r in rows]) if (headers or rows) else 0
    if ncol == 0:
        return
    tbl = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=ncol)
    set_table_grid(tbl)
    colw = Inches(CONTENT_W / ncol)
    ri = 0
    if headers:
        for j in range(ncol):
            c = tbl.rows[0].cells[j]; c.width = colw
            shade(c, TEAL)
            fmt_cell(c, headers[j] if j < len(headers) else "", color=WHITE, bold=True, size=9,
                     align=("left" if j == 0 else "center"))
        ri = 1
    for i, row in enumerate(rows):
        is_emph = i in emph
        for j in range(ncol):
            c = tbl.rows[ri + i].cells[j]; c.width = colw
            if is_emph: shade(c, EMPH)
            fmt_cell(c, row[j] if j < len(row) else "", color=BLACK, bold=is_emph, size=9,
                     align=("left" if j == 0 else "center"))
    if blk.get("source"):
        add_source(doc, blk["source"])

def add_image(doc, blk):
    """Paste a pre-built exhibit/chart image (house convention for complex exhibits),
    full content-width by default, centered, with an optional gold italic source line."""
    path = blk.get("path")
    if not path or not os.path.exists(path):
        sys.stderr.write("WARN: image not found %r (skipped)\n" % path)
        return
    width = blk.get("width_in", CONTENT_W)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))
    if blk.get("source"):
        add_source(doc, blk["source"])

def add_source(doc, text):
    p = doc.add_paragraph(style=S_BODY)
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)
    for r in p.runs:
        r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GOLD

# --- main build ------------------------------------------------------------------
def build(spec, out_path):
    if not os.path.exists(TEMPLATE):
        sys.exit("ERROR: golden template not found at %s" % TEMPLATE)
    doc = Document(TEMPLATE)

    title = (spec.get("title") or "").strip()

    # House style (golden HARVEST summary) is a CLEAN titled document:
    #   logo header (from template) -> title -> first H1 (Executive Summary).
    # The Addepar/Form/Radiant deck cover (hero banner -> gold "INVESTMENT SUMMARY"
    # label -> 4-column facts table) is now OPT-IN: set "cover": true in the spec.
    cover = bool(spec.get("cover"))
    if cover:
        add_banner(doc, title or "INVESTMENT SUMMARY", spec.get("banner_image"))
    if title:
        add_runs(doc.add_paragraph(style=S_TITLE), title)
    if cover:
        add_label(doc, spec.get("subtitle", "INVESTMENT SUMMARY"))
        if spec.get("facts") is not None:
            add_facts_table(doc, spec.get("facts") or {})

    # Body blocks
    for blk in spec.get("blocks", []):
        t = blk.get("type")
        if t == "h1":
            add_runs(doc.add_paragraph(style=S_H1), blk["text"])
        elif t == "h2":
            add_runs(doc.add_paragraph(style=S_H2), blk["text"])
        elif t == "p":
            add_runs(doc.add_paragraph(style=S_BODY), blk["text"])
        elif t == "bullets":
            for item in blk.get("items", []):
                p = doc.add_paragraph(style=S_BODY)
                add_runs(p, item["text"])
                set_bullet(p, item.get("level", 0))
        elif t == "table":
            add_data_table(doc, blk)
        elif t == "image":
            add_image(doc, blk)
        elif t == "source":
            add_source(doc, blk["text"])
        elif t == "risk_factors":
            add_runs(doc.add_paragraph(style=S_H1), "Risk Factors")
            for line in RISK_FACTORS:
                p = doc.add_paragraph(style=S_BODY); add_runs(p, line); set_bullet(p, 0)
        elif t == "disclaimer":
            add_runs(doc.add_paragraph(style=S_H1), "Disclaimer")
            lines = {"short": DISCLAIMER_SHORT, "long": DISCLAIMER_LONG}.get(
                blk.get("variant"), DISCLAIMER_STANDARD)
            for line in lines:
                p = doc.add_paragraph(style=S_BODY); add_runs(p, line); set_bullet(p, 0)
        else:
            sys.stderr.write("WARN: unknown block type %r (skipped)\n" % t)

    if spec.get("canadian"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("For Canadian residents only"); r.font.name = "Arial"; r.font.size = Pt(8)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path

def main():
    ap = argparse.ArgumentParser(description="Build an August Group Investment Summary .docx from a JSON spec.")
    ap.add_argument("spec", help="path to the JSON spec file")
    ap.add_argument("--out", help="explicit output .docx path (overrides spec filename/output_dir)")
    args = ap.parse_args()

    with open(args.spec, "r", encoding="utf-8") as f:
        spec = json.load(f)

    if args.out:
        out_path = args.out
    else:
        fname = spec.get("filename") or ((spec.get("title") or "Investment Summary").title() + " - Investment Summary.docx")
        if not fname.lower().endswith(".docx"):
            fname += ".docx"
        out_path = os.path.join(spec.get("output_dir", "."), fname)

    path = build(spec, out_path)
    print("SAVED:", os.path.abspath(path))

if __name__ == "__main__":
    main()
